# app.py
# Run:
#   pip install fastapi uvicorn pdfplumber groq python-multipart cryptography
#   python -m uvicorn app:app --port 8000

import json, os, io, re
from datetime  import datetime, timezone
from typing    import Optional

import pdfplumber

# DOCX support — pip install python-docx
try:
    import docx as python_docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
from fastapi              import FastAPI, UploadFile, File, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic             import BaseModel
from groq                 import Groq, RateLimitError

import db
from key_vault import (
    decrypt_client_groq_key, encrypt_groq_key, decrypt_groq_key,
    make_sentinel_token, verify_sentinel_token
)
from mail_chat import (
    llm, safe_json, FAST_MODEL, SMART_MODEL,
    MailRequest, MailGenRequest, ToneCheckRequest, ChatRequest,
    JobDescriptionRequest, SentinelRequest,
    handle_mail_generate, handle_tone_check, handle_mail_send,
    handle_hr_log, handle_chat, handle_chat_clear, handle_chat_history,
    handle_mail_details
)

app = FastAPI(title="SentinelX Server")

ALLOWED_ORIGINS = os.environ.get("ALLOWED_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_methods=["*"],
    allow_headers=["*"],
)

GROQ_KEY_TTL = 8 * 3600   # 8 hours

# ── Auth dependency ───────────────────────────────────────────────────────────

def get_groq_key(
    sentinel_x:    Optional[str] = Header(None, alias="X-Sentinel-Id"),
    authorization: Optional[str] = Header(None),
) -> Optional[str]:
    """
    Bearer token auth: Authorization: Bearer <sentinel_token>
    User identity:     X-Sentinel-Id: <sentinelX>
    Verifies HMAC token, fetches triple-AES-encrypted groq_key from Redis, decrypts.
    Returns None -> caller falls back to server GROQ_API_KEY.
    """
    if not (authorization and sentinel_x):
        return None
    if not authorization.startswith("Bearer "):
        return None
    token = authorization.removeprefix("Bearer ").strip()
    if not verify_sentinel_token(token, sentinel_x):
        return None
    encrypted = db.redis_client().get(f"gk:{sentinel_x}")
    if not encrypted:
        return None
    if isinstance(encrypted, bytes):
        encrypted = encrypted.decode()
    return decrypt_groq_key(encrypted, sentinel_x)


# ── Resume schema ─────────────────────────────────────────────────────────────

RESUME_SCHEMA = """{
  "basic_info": { "name":"","email":"","phone":"","location":"","linkedin":"","github":"","portfolio":"" },
  "skills": { "technical":[],"soft":[],"languages":[],"tools":[] },
  "professional_certifications": [ {"name":"","issuer":"","year":""} ],
  "experience": [ {"company":"","title":"","location":"","start_date":"","end_date":"","current":false,"responsibilities":[]} ],
  "projects": [ {"name":"","description":[],"tech_stack":[],"url":""} ],
  "education": [ {"institution":"","degree":"","field":"","start_date":"","end_date":"","grade":""} ]
}"""

PARSE_SYSTEM = (
    "You are a resume parser. Extract structured data verbatim.\n"
    "Return ONLY valid JSON — no markdown, no backticks.\n"
    "Use exactly this schema, empty string/array for unknowns:\n"
    + RESUME_SCHEMA
)

# ── Routes ────────────────────────────────────────────────────────────────────

@app.post("/clear_groq_key")
def clear_groq_key(
    sentinel_x: str = Header(..., alias="X-Sentinel-Id"),
):
    db.redis_client().delete(f"gk:{sentinel_x}")
    return {"status": "ok"}

@app.get("/health")
def health(
    sentinel_x:         str           = Header(...,  alias="X-Sentinel-Id"),
    crypt_groq_key: Optional[str] = Header(None, alias="X-Groq-Key"),
):
    """
    Health check + groq_key registration.

    First call (key registration):
      X-Sentinel-Id: <sentinelX>
      X-Groq-Key:    <server-AES-GCM encrypted groq_key from popup.js encryptForServer()>

    Subsequent calls (token refresh / status check):
      X-Sentinel-Id: <sentinelX>
      Authorization: Bearer <sentinel_token>   (no X-Groq-Key needed)

    Response:
      sentinel_token  -- HMAC bearer token (stable, no expiry in token itself)
      groq_key_valid  -- true if client key is active in Redis
    """
    result   = db.get_resume(sentinel_x)
    response = {
        "status":         "ok",
        "resume_present": result.get("resume_present", False),
        "last_updated":   result.get("last_updated"),
        "sentinel_token": None,
        "groq_key_valid": False,
    }
    if crypt_groq_key:
        groq_key = decrypt_client_groq_key(crypt_groq_key, sentinel_x)
        if groq_key:
            try:
                Groq(api_key=groq_key).chat.completions.create(
                    model="llama-3.1-8b-instant",
                    max_tokens=1,
                    messages=[{"role": "user", "content": "hi"}],
                )
                db.redis_client().setex(
                    f"gk:{sentinel_x}",
                    GROQ_KEY_TTL,
                    encrypt_groq_key(groq_key, sentinel_x)
                )
                response["sentinel_token"] = make_sentinel_token(sentinel_x)
                response["groq_key_valid"] = True
            except Exception:
                response["groq_key_valid"] = False
        return response

    # No key upload — check if one already lives in Redis
    if db.redis_client().exists(f"gk:{sentinel_x}"):
        response["sentinel_token"] = make_sentinel_token(sentinel_x)
        response["groq_key_valid"] = True
    return response


# ── Resume validation helpers ─────────────────────────────────────────────────

# Keywords that strongly indicate resume content
_RESUME_SIGNALS = [
    "experience", "education", "skills", "work", "employment", "university",
    "college", "degree", "bachelor", "master", "phd", "engineer", "developer",
    "manager", "intern", "internship", "project", "certification", "resume",
    "cv", "curriculum vitae", "responsibilities", "achievement", "linkedin",
    "github", "portfolio", "objective", "summary", "profile",
]

def _looks_like_resume(text: str) -> bool:
    """Return True if text has enough resume-like signals."""
    lower = text.lower()
    hits  = sum(1 for kw in _RESUME_SIGNALS if kw in lower)
    # Require at least 4 distinct signals and minimum 150 chars of content
    return hits >= 4 and len(text.strip()) >= 150


def _extract_pdf_text(contents: bytes) -> str:
    pages = []
    with pdfplumber.open(io.BytesIO(contents)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
    return "\n".join(pages).strip()


def _extract_docx_text(contents: bytes) -> str:
    if not DOCX_AVAILABLE:
        raise HTTPException(
            400,
            "DOCX support unavailable on this server. "
            "Please upload a PDF instead, or ask the admin to install python-docx."
        )
    doc   = python_docx.Document(io.BytesIO(contents))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text from tables (skills tables etc.)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paras.append(cell.text.strip())
    return "\n".join(paras).strip()


@app.post("/resume/upload")
async def upload_resume(
    sentinel_x: str           = Header(...,  alias="X-Sentinel-Id"),
    file:       UploadFile    = File(...),
    groq_key:   Optional[str] = Depends(get_groq_key),
):
    ext      = (file.filename or "").rsplit(".", 1)[-1].lower()
    contents = await file.read()

    # ── Text-based formats: PDF and DOCX ─────────────────────────────────────
    if ext in ("pdf", "docx"):
        try:
            if ext == "pdf":
                raw_text = _extract_pdf_text(contents)
            else:
                raw_text = _extract_docx_text(contents)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(422, f"Could not read file: {e}")

        if not raw_text:
            raise HTTPException(422, f"No text could be extracted from the {ext.upper()}. "
                                     "Try saving it as a text-based PDF or DOCX (not a scanned image).")

        # ── Resume content validation ─────────────────────────────────────────
        if not _looks_like_resume(raw_text):
            raise HTTPException(
                422,
                "This file does not appear to contain a resume. "
                "Please upload your CV or resume — the document must include sections like "
                "experience, education, and skills."
            )

        # ── LLM parse with retry on malformed output ──────────────────────────
        last_err = None
        structured = None
        for attempt in range(2):          # up to 2 attempts
            try:
                raw_ai = llm(PARSE_SYSTEM, raw_text[:12000], 4000,
                             groq_key=groq_key, model=SMART_MODEL)
                structured = safe_json(raw_ai)
                break
            except HTTPException as e:
                last_err = e
                if e.status_code != 500:  # don't retry auth/rate errors
                    raise

        if structured is None:
            raise last_err or HTTPException(500, "Resume parsing failed after retries. Please try again.")

    # ── JSON format ───────────────────────────────────────────────────────────
    elif ext == "json":
        try:
            structured = json.loads(contents.decode("utf-8"))
        except json.JSONDecodeError:
            raise HTTPException(400, "Invalid JSON file.")
        required = ["basic_info", "skills", "experience", "projects", "education"]
        missing  = [k for k in required if k not in structured]
        if missing:
            raise HTTPException(400, f"Missing required keys: {missing}")
        for alt in ["professional Certification", "professionalCertification"]:
            if alt in structured and "professional_certifications" not in structured:
                structured["professional_certifications"] = structured.pop(alt)
        structured.setdefault("professional_certifications", [])

    else:
        raise HTTPException(400, "Unsupported file type. Upload a PDF, DOCX, or JSON resume.")

    return db.save_resume(sentinel_x, structured)


@app.get("/resume/get")
def get_resume(sentinel_x: str = Header(..., alias="X-Sentinel-Id")):
    result = db.get_resume(sentinel_x)
    if result.get("resume_present"):
        return {**result, "status": "retrieved"}
    return {"status": "not_found", "resume_present": False, "last_updated": None, "resume": "NO_RESUME_FOUND"}


@app.delete("/resume/clear")
def clear_resume(sentinel_x: str = Header(..., alias="X-Sentinel-Id")):
    return db.clear_resume(sentinel_x)


# ── Job analysis ──────────────────────────────────────────────────────────────

class JobAnalyseRequest(BaseModel):
    job:    dict
    resume: dict


@app.post("/job/analyse")
def analyse_job(req: JobAnalyseRequest, groq_key: Optional[str] = Depends(get_groq_key)):
    system = (
        "You are an expert job application assistant.\n"
        "Return ONLY valid JSON — no markdown, no backticks.\nSchema:\n"
        "Do not ignore experience pick all skills from description and match with resume skills. Do not hallucinate skills that are not in the description.\n"
        '{"fit_score":<1-10>,"fit_reason":"","cover_letter":"write a cover letter based on the job description and resume","key_skills_match":[],"missing_skills":[],"apply_recommendation":true}\n\n'
        f"CANDIDATE RESUME:\n{json.dumps(req.resume, indent=2)}"
    )
    user = (
        f"Analyse this job posting.\n\nTitle: {req.job.get('title','')}\n"
        f"Company: {req.job.get('company','')}\nLocation: {req.job.get('location','')}\n"
        f"Description:\n{req.job.get('description','')}"
    )
    try:
        return safe_json(llm(system, user, 5000, groq_key=groq_key, model=SMART_MODEL))
    except Exception as e:
        raise HTTPException(500, f"LLM error: {e}")


# ── Mail + Chat ───────────────────────────────────────────────────────────────

@app.post("/mail/generate")
def mail_generate(req: MailGenRequest, groq_key: Optional[str] = Depends(get_groq_key)):
    return handle_mail_generate(req, groq_key)

@app.post("/mail/tone-check")
def mail_tone_check(req: ToneCheckRequest, groq_key: Optional[str] = Depends(get_groq_key)):
    return handle_tone_check(req, groq_key)

@app.post("/mail/send")
def mail_send(req: MailRequest):
    return handle_mail_send(req)

@app.get("/mail/hr_log")
def hr_log(sentinel_x: str = Header(..., alias="X-Sentinel-Id")):
    return handle_hr_log(sentinel_x)

@app.post("/generate/mail_details")
def generate_mail_details(req: JobDescriptionRequest, groq_key: Optional[str] = Depends(get_groq_key)):
    return handle_mail_details(req, groq_key)

@app.post("/chat")
def chat(req: ChatRequest, groq_key: Optional[str] = Depends(get_groq_key)):
    return handle_chat(req, groq_key)

@app.get("/chat/history")
def chat_history(sentinel_x: str = Header(..., alias="X-Sentinel-Id")):
    return handle_chat_history(sentinel_x)

@app.delete("/chat/clear")
def chat_clear(sentinel_x: str = Header(..., alias="X-Sentinel-Id")):
    return handle_chat_clear(sentinel_x)


# ── Identity ──────────────────────────────────────────────────────────────────

@app.post("/identity/register")
def register_identity(req: SentinelRequest):
    if not req.sentinel_x:
        raise HTTPException(400, "sentinel_x required")
    result = db.register_user(req.sentinel_x)
    return {"status": "ok", "user_status": result.get("status")}

@app.post("/identity/verify")
def verify_identity(req: SentinelRequest):
    if not req.sentinel_x:
        raise HTTPException(400, "sentinel_x required")
    valid = db.verify_user(req.sentinel_x)
    if valid:
        db.touch_user(req.sentinel_x)
    return {"valid": valid}
